"""
generate_report.py
==================
Reusable script: converts any db_document.json → professional Excel + DOCX report.

Usage:
    python generate_report.py <path/to/db_document.json>
    python generate_report.py db_document.json          # defaults to current dir

Outputs:
    output/<company_name>/
        ├── <company_name>_report.xlsx
        └── <company_name>_report.docx
"""

import sys, os, re, json, io, textwrap, tempfile
from pathlib import Path
from datetime import datetime

import requests
from PIL import Image as PILImage

import openpyxl
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                              GradientFill)
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

# ══════════════════════════════════════════════════════════════════════════════
# CONFIG
# ══════════════════════════════════════════════════════════════════════════════

LOGO_DEV_KEY   = "pk_DOneOcGwSAau_ztcDSzJYw"
LOGO_SIZE_PX   = 80          # px we'll request from logo.dev
FONT_NAME      = "Calibri"

# Brand palette
C_DARK         = "0D1F3C"    # deep navy
C_ACCENT       = "1A6BCC"    # corporate blue
C_ACCENT_LIGHT = "D6E8FF"    # pale blue fill
C_GREEN        = "1A7A4A"    # success green
C_GREEN_LIGHT  = "D4EDDA"
C_AMBER        = "B76E00"    # warning amber
C_AMBER_LIGHT  = "FFF3CD"
C_RED          = "B71C1C"    # danger red
C_RED_LIGHT    = "FDECEA"
C_GREY_BG      = "F5F7FA"    # alternating row
C_GREY_LINE    = "DEE2E8"    # border colour
C_WHITE        = "FFFFFF"
C_TEXT_DARK    = "1A1A2E"
C_TEXT_MUTED   = "6B7280"

# Severity → colour mapping (Excel and DOCX)
SEV_PALETTE = {
    "critical": (C_RED,      C_RED_LIGHT),
    "high":     (C_AMBER,    C_AMBER_LIGHT),
    "medium":   ("7B5800",   "FFFBEA"),
    "low":      (C_GREEN,    C_GREEN_LIGHT),
    "p0":       (C_RED,      C_RED_LIGHT),
    "p1":       (C_AMBER,    C_AMBER_LIGHT),
    "p2":       ("7B5800",   "FFFBEA"),
    "p3":       (C_GREEN,    C_GREEN_LIGHT),
}

# ══════════════════════════════════════════════════════════════════════════════
# UTILITY — safe value extraction, cleaning, formatting
# ══════════════════════════════════════════════════════════════════════════════

_NULL_VALS = {None, "", "null", "none", "n/a", "na", "—", "unable to verify",
              "not available", "not found", "undefined"}

def is_empty(v):
    """Return True when a value carries no useful information."""
    if v is None:
        return True
    if isinstance(v, str) and v.strip().lower() in _NULL_VALS:
        return True
    if isinstance(v, (list, dict)) and len(v) == 0:
        return True
    return False

def safe(v, fallback="—"):
    """Scalar value → clean string, empty → fallback."""
    if is_empty(v):
        return fallback
    if isinstance(v, bool):
        return "Yes" if v else "No"
    if isinstance(v, float):
        return f"{v:,.4g}"
    if isinstance(v, list):
        parts = [safe(i) for i in v if not is_empty(i)]
        return ", ".join(parts) if parts else fallback
    return str(v).strip()

def clean_text(v):
    """Remove internal JSON artefacts from long text fields."""
    s = safe(v)
    if s == "—":
        return s
    # Remove leftover raw URLs that are wrapped in markdown [text](url)
    s = re.sub(r'\[([^\]]+)\]\(https?://[^\)]+\)', r'\1', s)
    # Collapse 3+ newlines
    s = re.sub(r'\n{3,}', '\n\n', s)
    # Remove leading/trailing whitespace per line
    lines = [l.rstrip() for l in s.split('\n')]
    return '\n'.join(lines).strip()

def is_url(v):
    return isinstance(v, str) and re.match(r'https?://', v.strip())

def fmt_date(v, fallback="—"):
    if is_empty(v):
        return fallback
    s = str(v)
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f",
                "%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%d", "%Y-%m-%dT%H:%M:%SZ"):
        try:
            return datetime.strptime(s[:26], fmt).strftime("%-d %b %Y")
        except Exception:
            pass
    return s[:10]

def fmt_num(v, suffix=""):
    if is_empty(v):
        return "—"
    try:
        return f"{int(float(str(v))):,}{suffix}"
    except Exception:
        return safe(v)

def star_bar(rating, max_stars=5):
    """Return a star-bar string, e.g. ★★★★☆ 4.3"""
    try:
        r = float(rating)
        full = int(r)
        half = 1 if (r - full) >= 0.5 else 0
        empty = max_stars - full - half
        bar = "★" * full + ("½" if half else "") + "☆" * empty
        return f"{bar}  {r:.1f}"
    except Exception:
        return safe(rating)

def clean_domain(domain_str):
    """https://univest.in/ → univest.in"""
    d = re.sub(r'^https?://', '', str(domain_str or "")).rstrip('/')
    return d

def slug(name):
    return re.sub(r'[^\w]+', '_', str(name).strip()).strip('_')

# ══════════════════════════════════════════════════════════════════════════════
# LOGO FETCH
# ══════════════════════════════════════════════════════════════════════════════

def fetch_logo_bytes(domain, company_name):
    """Try logo.dev first (domain, then name); fall back to initials PNG."""
    domain_clean = clean_domain(domain)
    for query in [domain_clean, re.sub(r'\s+', '', company_name.lower())]:
        try:
            url = f"https://img.logo.dev/{query}?token={LOGO_DEV_KEY}&size={LOGO_SIZE_PX}&format=png"
            r = requests.get(url, timeout=6, headers={"User-Agent": "Mozilla/5.0"})
            if r.status_code == 200 and r.headers.get("content-type", "").startswith("image"):
                return r.content, "png"
        except Exception:
            pass
    # Fallback: generate an initials badge with Pillow
    return _make_initials_logo(company_name), "png"

def _make_initials_logo(name):
    """Create a coloured square with company initials using Pillow."""
    from PIL import ImageDraw, ImageFont
    size = LOGO_SIZE_PX * 2          # render at 2x, downscale
    img  = PILImage.new("RGB", (size, size), color=(26, 107, 204))
    draw = ImageDraw.Draw(img)
    initials = "".join(w[0].upper() for w in name.split()[:2])
    # Simple font sizing
    font_size = size // 2
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), initials, font=font)
    x = (size - (bbox[2] - bbox[0])) // 2 - bbox[0]
    y = (size - (bbox[3] - bbox[1])) // 2 - bbox[1]
    draw.text((x, y), initials, fill=(255, 255, 255), font=font)
    buf = io.BytesIO()
    img = img.resize((LOGO_SIZE_PX, LOGO_SIZE_PX), PILImage.LANCZOS)
    img.save(buf, format="PNG")
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# DATA EXTRACTOR — normalises every section from JSON
# ══════════════════════════════════════════════════════════════════════════════

class ReportData:
    """Parses db_document.json and exposes clean, typed attributes."""

    def __init__(self, path):
        with open(path, encoding="utf-8") as f:
            raw = json.load(f)

        self.project_name  = safe(raw.get("project_name"))
        self.domain        = safe(raw.get("domain"))
        self.ingestion_date = fmt_date(raw.get("ingestion_date"))

        src = raw.get("data_sources", {})

        # ── Company Profile ──────────────────────────────────────────────────
        cp = src.get("company_profile", {}).get("data", {})
        self.company_name       = safe(cp.get("company_name"), self.project_name)
        self.company_domain     = safe(cp.get("domain"), self.domain)
        self.playstore_link     = safe(cp.get("playstore_link"))
        self.appstore_link      = safe(cp.get("appstore_link"))
        self.youtube_channel    = safe(cp.get("youtube_official_channel"))
        self.linkedin_page      = safe(cp.get("linkedin_company_page"))
        self.year_founded       = safe(cp.get("year_founded"))
        self.hq_location        = safe(cp.get("exact_hq_location"))
        self.locations          = self._flatten_list(cp.get("locations_operating_in", []))
        self.industry           = clean_text(cp.get("industry_and_segment"))
        self.platforms          = safe(cp.get("available_platforms"))
        self.employee_count     = safe(cp.get("employee_count"))
        self.funding_raised     = safe(cp.get("funding_raised"))
        self.funding_stage      = safe(cp.get("funding_stage"))
        self.no_of_users        = safe(cp.get("no_of_users"))
        self.annual_revenue     = safe(cp.get("annual_revenue"))
        self.key_positioning    = clean_text(cp.get("key_positioning"))
        self.revenue_model      = safe(cp.get("revenue_model"))
        self.pricing_tiers      = self._flatten_list(cp.get("pricing_tiers", []))
        self.target_segments    = self._flatten_list(cp.get("target_customer_segments", []))
        self.tech_stack         = self._flatten_list(cp.get("tech_stack_highlights", []))
        self.milestones         = self._flatten_list(cp.get("milestones", []))
        self.new_features       = self._flatten_list(cp.get("new_features_launched", []))
        self.founders           = self._flatten_list(cp.get("names_of_founders", []))
        self.csuite             = self._flatten_list(cp.get("c-suite_officer", cp.get("c_suite_officer", [])))
        self.competitors        = self._parse_structured_list(cp.get("competitors", []))
        self.partnerships       = self._parse_structured_list(cp.get("recent_partnerships_and_integrations", []))
        self.strategic_moves    = self._parse_structured_list(cp.get("strategic_moves", []))
        self.differentiators    = self._parse_structured_list(cp.get("differentiators", []))
        self.user_complaints    = self._parse_structured_list(cp.get("user_complaints", []))
        self.regulatory         = self._parse_structured_list(cp.get("regulatory_and_legal_issues", []))
        self.current_problems   = self._parse_structured_list(cp.get("current_problems_struggling_with", []))
        self.other_details      = self._flatten_list(cp.get("other_crucial_details", []))
        self.market_sentiment   = cp.get("market_sentiment") or {}

        # ── Play Store ───────────────────────────────────────────────────────
        ps_raw = src.get("play_store", {})
        ps_meta = ps_raw.get("extracted_data", {}).get("metadata", {})
        ps_ra   = ps_raw.get("extracted_data", {}).get("review_analysis", {})
        ps_rev  = ps_raw.get("extracted_data", {}).get("reviews", [])
        self.ps = {
            "title":          safe(ps_meta.get("title")),
            "developer":      safe(ps_meta.get("developer")),
            "score":          safe(ps_meta.get("score")),
            "ratings":        fmt_num(ps_meta.get("ratings")),
            "reviews":        fmt_num(ps_meta.get("reviews")),
            "installs":       safe(ps_meta.get("installs")),
            "genre":          safe(ps_meta.get("genre")),
            "released":       safe(ps_meta.get("released")),
            "version":        safe(ps_meta.get("version")),
            "content_rating": safe(ps_meta.get("content_rating")),
            "free":           "Free" if ps_meta.get("free") else "Paid",
            "url":            safe(ps_meta.get("url")),
            "privacy_policy": safe(ps_meta.get("developer_privacy_policy")),
            "min_android":    safe(ps_meta.get("min_android_version")),
            "total_reviews":  fmt_num(ps_ra.get("total_reviews")),
            "avg_rating":     safe(ps_ra.get("average_rating")),
            "latest_review":  fmt_date(ps_ra.get("latest_review_date")),
            "oldest_review":  fmt_date(ps_ra.get("oldest_review_date")),
            "rating_dist":    ps_ra.get("rating_distribution", {}),
            "reviews_list":   self._parse_ps_reviews(ps_rev),
        }

        # ── App Store ────────────────────────────────────────────────────────
        as_raw  = src.get("app_store", {})
        as_meta = as_raw.get("extracted_data", {}).get("metadata", {})
        self.app = {
            "title":         safe(as_meta.get("trackName")),
            "developer":     safe(as_meta.get("artistName")),
            "score":         safe(as_meta.get("averageUserRating")),
            "ratings":       fmt_num(as_meta.get("userRatingCount")),
            "version":       safe(as_meta.get("version")),
            "min_ios":       safe(as_meta.get("minimumOsVersion")),
            "genre":         safe(as_meta.get("primaryGenreName")),
            "released":      fmt_date(as_meta.get("releaseDate")),
            "price":         safe(as_meta.get("formattedPrice")),
            "url":           safe(as_meta.get("trackViewUrl")),
            "content_rating":safe(as_meta.get("contentAdvisoryRating")),
        }

        # ── Transcripts ──────────────────────────────────────────────────────
        tr = src.get("internal_transcripts", {})
        self.transcript = {
            "source_file":   safe(tr.get("source_file")),
            "total_signals": safe(tr.get("total_signals")),
            "classifier":    safe(tr.get("classifier_used")),
            "meeting_type":  safe((tr.get("metadata") or {}).get("meeting_type")),
            "processed_at":  fmt_date((tr.get("metadata") or {}).get("processed_at")),
            "signals":       self._parse_signals(tr.get("signals", [])),
        }

        # ── Reddit ───────────────────────────────────────────────────────────
        reddit_raw = src.get("reddit", {})
        self.reddit_posts = []
        for block in reddit_raw.values():
            for post in block.get("posts", []):
                if not is_empty(post.get("title")):
                    self.reddit_posts.append({
                        "title":        clean_text(post.get("title")),
                        "subreddit":    safe(post.get("subreddit")),
                        "author":       safe(post.get("author")),
                        "score":        fmt_num(post.get("score")),
                        "comments":     fmt_num(post.get("num_comments")),
                        "url":          safe(post.get("url")),
                        "selftext":     clean_text(post.get("selftext")),
                    })

        # ── YouTube ──────────────────────────────────────────────────────────
        yt_raw = src.get("youtube", {})
        self.youtube_videos = []
        for block in yt_raw.values():
            for vid in (block if isinstance(block, list) else []):
                if not is_empty(vid.get("title")):
                    self.youtube_videos.append({
                        "title":       clean_text(vid.get("title")),
                        "url":         safe(vid.get("url")),
                        "video_id":    safe(vid.get("video_id")),
                        "description": clean_text(vid.get("description")),
                        "views":       fmt_num(vid.get("view_count")),
                        "likes":       fmt_num(vid.get("like_count")),
                        "published":   fmt_date(vid.get("published_at")),
                        "scraped_at":  fmt_date(vid.get("scraped_at")),
                    })

        # ── Agent 2 — Problems ───────────────────────────────────────────────
        ag2 = raw.get("agent2_output", {})
        self.problems = []
        for p in ag2.get("problems", []):
            self.problems.append({
                "id":               safe(p.get("problem_id")),
                "problem":          clean_text(p.get("problem")),
                "severity":         safe(p.get("severity")),
                "frequency":        safe(p.get("frequency")),
                "category":         safe(p.get("category")),
                "user_type":        safe(p.get("user_type")),
                "sources":          self._flatten_list(p.get("source_mix", [])),
                "competitor_issue": safe(p.get("competitor_has_same_issue")),
                "evidence":         [clean_text(e) for e in (p.get("evidence") or []) if not is_empty(e)],
            })
        self.total_problems     = safe(ag2.get("total_problems"))
        self.top_categories     = self._flatten_list(ag2.get("top_categories", []))
        self.high_severity_count= safe(ag2.get("high_severity_count"))

        # ── Agent 3 — Insights ───────────────────────────────────────────────
        ag3 = raw.get("agent3_output", {})
        self.insights = []
        for i in ag3.get("insights", []):
            self.insights.append({
                "id":              safe(i.get("insight_id")),
                "insight":         clean_text(i.get("insight")),
                "priority":        safe(i.get("priority")),
                "confidence":      safe(i.get("confidence")),
                "theme":           safe(i.get("theme")),
                "root_cause":      clean_text(i.get("root_cause")),
                "evidence":        clean_text(i.get("evidence_summary")),
                "competitor_gap":  clean_text(i.get("competitor_gap")),
                "opportunity":     clean_text(i.get("opportunity_size")),
                "implication":     clean_text(i.get("implication")),
                "support_ids":     self._flatten_list(i.get("supporting_problem_ids", [])),
            })
        self.total_insights     = safe(ag3.get("total_insights"))
        self.critical_count     = safe(ag3.get("critical_count"))
        self.dominant_theme     = safe(ag3.get("dominant_theme"))
        self.strategic_risk     = clean_text(ag3.get("key_strategic_risk"))
        self.biggest_opp        = clean_text(ag3.get("biggest_opportunity"))

        # ── Agent 4 — Briefs ─────────────────────────────────────────────────
        ag4 = raw.get("agent4_output", {})
        self.briefs = []
        for b in ag4.get("briefs", []):
            self.briefs.append({
                "id":           safe(b.get("brief_id")),
                "feature":      safe(b.get("feature_name")),
                "priority":     safe(b.get("priority")),
                "effort":       safe(b.get("effort")),
                "insight_ref":  safe(b.get("addresses_insight")),
                "problem":      clean_text(b.get("problem")),
                "why_now":      clean_text(b.get("why_now")),
                "solution":     clean_text(b.get("solution")),
                "impact":       clean_text(b.get("expected_impact")),
                "metric":       clean_text(b.get("success_metric")),
                "user_flow":    [clean_text(s) for s in (b.get("user_flow") or []) if not is_empty(s)],
            })
        self.total_briefs       = safe(ag4.get("total_briefs"))
        self.sprint_focus       = clean_text(ag4.get("recommended_sprint_focus"))

    # ── helpers ──────────────────────────────────────────────────────────────

    def _flatten_list(self, lst):
        out = []
        for item in (lst or []):
            if is_empty(item):
                continue
            if isinstance(item, str):
                out.append(item.strip())
            elif isinstance(item, dict):
                # Pick the most descriptive text key
                for key in ("description", "feature", "move", "issue", "detail", "note", "text"):
                    if not is_empty(item.get(key)):
                        out.append(clean_text(item[key]))
                        break
                else:
                    # Fall back to joining all non-url values
                    parts = [str(v) for v in item.values()
                             if not is_empty(v) and not is_url(str(v))]
                    if parts:
                        out.append(" | ".join(parts))
            else:
                out.append(str(item))
        return out

    def _parse_structured_list(self, lst):
        """Return list of dicts, skip entries where all useful fields are empty."""
        out = []
        for item in (lst or []):
            if is_empty(item):
                continue
            if isinstance(item, str):
                out.append({"value": item})
            elif isinstance(item, dict):
                cleaned = {k: v for k, v in item.items()
                           if not is_empty(v) and k not in ("effect",)}
                if cleaned:
                    out.append(cleaned)
        return out

    def _parse_ps_reviews(self, lst):
        out = []
        for rev in (lst or []):
            content = clean_text(rev.get("content") or rev.get("text"))
            if is_empty(content):
                continue
            out.append({
                "author":  safe(rev.get("author") or rev.get("userName")),
                "rating":  safe(rev.get("rating") or rev.get("score")),
                "content": content,
                "date":    fmt_date(rev.get("date") or rev.get("at")),
                "reply":   clean_text(rev.get("reply_text") or rev.get("replyText")),
                "version": safe(rev.get("version")),
            })
        return out

    def _parse_signals(self, lst):
        out = []
        for s in (lst or []):
            content = clean_text(s.get("content"))
            if is_empty(content):
                continue
            out.append({
                "id":         safe(s.get("signal_id")),
                "type":       safe(s.get("signal_type")),
                "confidence": s.get("confidence"),
                "content":    content,
            })
        return out


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL BUILDER
# ══════════════════════════════════════════════════════════════════════════════

class ExcelBuilder:
    """Builds a polished, professional .xlsx workbook from ReportData."""

    # ── style constants ──────────────────────────────────────────────────────
    THIN   = Side(style="thin",   color=C_GREY_LINE)
    THICK  = Side(style="medium", color=C_ACCENT)
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

    def _F(self, bold=False, size=10, color=C_TEXT_DARK, italic=False):
        return Font(name=FONT_NAME, bold=bold, size=size,
                    color=color, italic=italic)

    def _P(self, color):
        return PatternFill("solid", fgColor=color)

    def _A(self, h="left", v="top", wrap=True):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    # ── cell helpers ─────────────────────────────────────────────────────────

    def _write(self, ws, row, col, value, *,
               bold=False, size=10, fg=C_TEXT_DARK, italic=False,
               fill=None, halign="left", wrap=True, border=True,
               link=None, num_fmt=None):
        c = ws.cell(row=row, column=col, value=value)
        c.font      = self._F(bold, size, fg, italic)
        c.alignment = self._A(halign, "top", wrap)
        if fill:
            c.fill = self._P(fill)
        if border:
            c.border = self.BORDER
        if link:
            c.hyperlink = link
            c.font = Font(name=FONT_NAME, size=size, color=C_ACCENT,
                          underline="single", bold=bold)
        if num_fmt:
            c.number_format = num_fmt
        return c

    def _header_row(self, ws, row, labels, widths=None, col_start=1):
        for i, lbl in enumerate(labels):
            c = ws.cell(row=row, column=col_start + i, value=lbl)
            c.font      = self._F(bold=True, size=10, color=C_WHITE)
            c.fill      = self._P(C_DARK)
            c.alignment = self._A("center", "center", False)
            c.border    = self.BORDER
        if widths:
            for i, w in enumerate(widths):
                ws.column_dimensions[get_column_letter(col_start + i)].width = w

    def _section_banner(self, ws, row, title, span, col_start=1):
        c = ws.cell(row=row, column=col_start, value=f"  {title}")
        c.font      = self._F(bold=True, size=11, color=C_WHITE)
        c.fill      = self._P(C_ACCENT)
        c.alignment = self._A("left", "center", False)
        c.border    = self.BORDER
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col_start,
                           end_row=row, end_column=col_start + span - 1)
        ws.row_dimensions[row].height = 22

    def _kv_row(self, ws, row, key, value, col=1, link=None):
        shade = C_GREY_BG if row % 2 == 0 else C_WHITE
        self._write(ws, row, col,   key,   bold=True, fill=shade)
        if link and is_url(str(value)):
            self._write(ws, row, col+1, str(value), fill=shade, link=str(value), wrap=True)
        else:
            self._write(ws, row, col+1, str(value), fill=shade, wrap=True)
        ws.row_dimensions[row].height = max(15, min(60, len(str(value)) // 3 + 15))

    def _sev_fill(self, sev_str):
        return SEV_PALETTE.get(sev_str.lower(), (C_TEXT_DARK, C_WHITE))[1]

    def _data_row(self, ws, row, values, shading=None):
        shade = shading or (C_GREY_BG if row % 2 == 0 else C_WHITE)
        for col, val in enumerate(values, 1):
            link = None
            if isinstance(val, tuple) and len(val) == 2:
                val, link = val
            self._write(ws, row, col, str(val) if val is not None else "—",
                        fill=shade, wrap=True, link=link)

    # ── sheet builders ───────────────────────────────────────────────────────

    def _sheet_overview(self, wb, d, logo_bytes, _logo_tmp_holder=None):
        ws = wb.create_sheet("Overview")
        ws.column_dimensions["A"].width = 26
        ws.column_dimensions["B"].width = 52
        ws.freeze_panes = "A3"

        # Logo + company name header block
        ws.row_dimensions[1].height = 70
        ws.merge_cells("A1:B1")
        title_cell = ws["A1"]
        title_cell.value    = f"  {d.company_name}  —  Intelligence Report"
        title_cell.font     = self._F(bold=True, size=16, color=C_WHITE)
        title_cell.fill     = self._P(C_DARK)
        title_cell.alignment= self._A("left", "center", False)

        if logo_bytes:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.write(logo_bytes); tmp.close()
            # Store path so build() can unlink AFTER wb.save()
            if _logo_tmp_holder is not None:
                _logo_tmp_holder.append(tmp.name)
            img = XLImage(tmp.name)
            img.width = 56; img.height = 56
            img.anchor = "A1"
            ws.add_image(img)

        ws.merge_cells("A2:B2")
        sub = ws["A2"]
        sub.value     = f"  {clean_domain(d.company_domain)}   ·   Generated {d.ingestion_date}"
        sub.font      = self._F(size=9, color="AAAAAA")
        sub.fill      = self._P(C_DARK)
        sub.alignment = self._A("left", "center", False)

        r = 3
        self._section_banner(ws, r, "Company Basics", 2); r += 1
        kv_pairs = [
            ("Company Name",      d.company_name,     False),
            ("Domain",            d.company_domain,   True),
            ("Year Founded",      d.year_founded,     False),
            ("HQ Location",       d.hq_location,      False),
            ("Operating In",      d.locations,        False),
            ("Industry",          d.industry,         False),
            ("Platforms",         d.platforms,        False),
            ("Employees",         d.employee_count,   False),
        ]
        for key, val, as_link in kv_pairs:
            if not is_empty(val):
                self._kv_row(ws, r, key, val, link=val if as_link else None); r += 1

        r += 1
        self._section_banner(ws, r, "Financials & Traction", 2); r += 1
        for key, val in [("Funding Stage", d.funding_stage), ("Funding Raised", d.funding_raised),
                         ("Annual Revenue", d.annual_revenue), ("No. of Users", d.no_of_users)]:
            if not is_empty(val):
                self._kv_row(ws, r, key, val); r += 1

        r += 1
        self._section_banner(ws, r, "Positioning", 2); r += 1
        for key, val in [("Key Positioning", d.key_positioning),
                         ("Revenue Model",   d.revenue_model)]:
            if not is_empty(val):
                self._kv_row(ws, r, key, val); r += 1

        if d.pricing_tiers:
            self._kv_row(ws, r, "Pricing Tiers", "\n".join(d.pricing_tiers)); r += 1
        if d.target_segments:
            self._kv_row(ws, r, "Target Segments", "  •  ".join(d.target_segments)); r += 1
        if d.tech_stack:
            self._kv_row(ws, r, "Tech Stack", ",  ".join(d.tech_stack)); r += 1

        r += 1
        self._section_banner(ws, r, "Official Links", 2); r += 1
        links = [("Play Store", d.playstore_link), ("App Store", d.appstore_link),
                 ("YouTube",    d.youtube_channel), ("LinkedIn",  d.linkedin_page)]
        for lbl, url in links:
            if not is_empty(url):
                self._kv_row(ws, r, lbl, url, link=url); r += 1

        r += 1
        self._section_banner(ws, r, "Market Sentiment", 2); r += 1
        ms = d.market_sentiment
        for key in ("overall", "analyst_view", "user_community_view"):
            val = ms.get(key)
            if not is_empty(val):
                self._kv_row(ws, r, key.replace("_", " ").title(), val); r += 1

        if d.csuite:
            r += 1
            self._section_banner(ws, r, "C-Suite Officers", 2); r += 1
            for officer in d.csuite:
                self._kv_row(ws, r, "", officer); r += 1

        if d.founders:
            r += 1
            self._section_banner(ws, r, "Founders", 2); r += 1
            self._kv_row(ws, r, "Founders", ",  ".join(d.founders)); r += 1

    def _sheet_app_store(self, wb, d):
        ws = wb.create_sheet("App Store")
        ws.column_dimensions["A"].width = 26
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 30
        ws.freeze_panes = "B2"

        ws.row_dimensions[1].height = 22
        self._section_banner(ws, 1, "App Store Comparison", 3)
        self._header_row(ws, 2, ["Metric", "🤖 Google Play Store", " Apple App Store"])

        rows = [
            ("App Title",      d.ps.get("title"),      d.app.get("title")),
            ("Rating",         star_bar(d.ps.get("avg_rating") or d.ps.get("score")),
                               star_bar(d.app.get("score"))),
            ("Total Ratings",  d.ps.get("ratings"),    d.app.get("ratings")),
            ("Total Reviews",  d.ps.get("reviews"),    "—"),
            ("Installs",       d.ps.get("installs"),   "—"),
            ("Genre",          d.ps.get("genre"),      d.app.get("genre")),
            ("Price",          d.ps.get("free"),       d.app.get("price")),
            ("Version",        d.ps.get("version"),    d.app.get("version")),
            ("Released",       d.ps.get("released"),   d.app.get("released")),
            ("Content Rating", d.ps.get("content_rating"), d.app.get("content_rating")),
            ("Min OS",         d.ps.get("min_android"), d.app.get("min_ios")),
            ("Store URL",      d.ps.get("url"),         d.app.get("url")),
            ("Privacy Policy", d.ps.get("privacy_policy"), "—"),
        ]
        r = 3
        for lbl, pv, av in rows:
            if is_empty(pv) and is_empty(av):
                continue
            shade = C_GREY_BG if r % 2 == 0 else C_WHITE
            self._write(ws, r, 1, lbl, bold=True, fill=shade)
            pv_s = safe(pv)
            if is_url(pv_s):
                self._write(ws, r, 2, pv_s, fill=shade, link=pv_s, wrap=True)
            else:
                self._write(ws, r, 2, pv_s, fill=shade, wrap=True)
            av_s = safe(av)
            if is_url(av_s):
                self._write(ws, r, 3, av_s, fill=shade, link=av_s, wrap=True)
            else:
                self._write(ws, r, 3, av_s, fill=shade, wrap=True)
            r += 1

        # Rating distribution
        rd = d.ps.get("rating_dist", {})
        if rd:
            r += 1
            self._section_banner(ws, r, "Play Store Rating Distribution", 3); r += 1
            self._header_row(ws, r, ["Stars", "Count", "Visual"]); r += 1
            total = sum(int(v) for v in rd.values() if str(v).isdigit()) or 1
            for star in ["5", "4", "3", "2", "1"]:
                cnt = int(rd.get(star, rd.get(int(star), 0)) or 0)
                bar = "█" * int((cnt / total) * 20)
                shade = C_GREY_BG if r % 2 == 0 else C_WHITE
                self._write(ws, r, 1, f"{'⭐' * int(star)} {star} stars", fill=shade)
                self._write(ws, r, 2, cnt, fill=shade, halign="center")
                self._write(ws, r, 3, bar, fill=shade,
                            fg=C_AMBER if star in ("1","2") else C_GREEN)
                r += 1

        ws.freeze_panes = "A3"

    def _sheet_reviews(self, wb, d):
        reviews = d.ps.get("reviews_list", [])
        if not reviews:
            return
        ws = wb.create_sheet("Reviews")
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 18
        ws.column_dimensions["C"].width = 8
        ws.column_dimensions["D"].width = 55
        ws.column_dimensions["E"].width = 30
        ws.column_dimensions["F"].width = 12
        ws.freeze_panes = "A2"

        self._header_row(ws, 1, ["#", "Author", "★", "Review", "Developer Reply", "Date"])
        for i, rev in enumerate(reviews, 1):
            r = i + 1
            shade = C_GREY_BG if i % 2 == 0 else C_WHITE
            sev_shade = {1: C_RED_LIGHT, 2: C_AMBER_LIGHT}.get(
                int(safe(rev["rating"]) or 3), shade)
            self._write(ws, r, 1, i,                fill=sev_shade, halign="center")
            self._write(ws, r, 2, rev["author"],    fill=sev_shade)
            self._write(ws, r, 3, rev["rating"],    fill=sev_shade, halign="center")
            self._write(ws, r, 4, rev["content"],   fill=sev_shade, wrap=True)
            reply = rev["reply"] if not is_empty(rev["reply"]) else "—"
            self._write(ws, r, 5, reply,            fill=sev_shade, wrap=True,
                        italic=True, fg=C_TEXT_MUTED)
            self._write(ws, r, 6, rev["date"],      fill=sev_shade)
            ws.row_dimensions[r].height = max(15, min(80, len(rev["content"]) // 4 + 15))

    def _sheet_social(self, wb, d):
        # ── Reddit ────────────────────────────────────────────────────────────
        if d.reddit_posts:
            ws = wb.create_sheet("Reddit")
            ws.column_dimensions["A"].width = 6
            ws.column_dimensions["B"].width = 48
            ws.column_dimensions["C"].width = 20
            ws.column_dimensions["D"].width = 10
            ws.column_dimensions["E"].width = 10
            ws.column_dimensions["F"].width = 18
            ws.freeze_panes = "A2"
            self._header_row(ws, 1, ["#", "Post Title", "Subreddit / Author", "Score", "Cmts", "Link"])
            for i, p in enumerate(d.reddit_posts, 1):
                r = i + 1
                shade = C_GREY_BG if i % 2 == 0 else C_WHITE
                self._write(ws, r, 1, i,                                      fill=shade, halign="center")
                self._write(ws, r, 2, p["title"],                             fill=shade, wrap=True)
                self._write(ws, r, 3, f"r/{p['subreddit']}  @{p['author']}", fill=shade)
                self._write(ws, r, 4, p["score"],                             fill=shade, halign="center")
                self._write(ws, r, 5, p["comments"],                          fill=shade, halign="center")
                url = p["url"]
                if is_url(url):
                    self._write(ws, r, 6, "Open ↗", fill=shade, link=url, halign="center")
                else:
                    self._write(ws, r, 6, "—", fill=shade, halign="center")

        # ── YouTube ───────────────────────────────────────────────────────────
        if d.youtube_videos:
            ws = wb.create_sheet("YouTube")
            ws.column_dimensions["A"].width = 6
            ws.column_dimensions["B"].width = 48
            ws.column_dimensions["C"].width = 12
            ws.column_dimensions["D"].width = 10
            ws.column_dimensions["E"].width = 14
            ws.freeze_panes = "A2"
            self._header_row(ws, 1, ["#", "Video Title", "Views", "Likes", "Link"])
            for i, v in enumerate(d.youtube_videos, 1):
                r = i + 1
                shade = C_GREY_BG if i % 2 == 0 else C_WHITE
                self._write(ws, r, 1, i,          fill=shade, halign="center")
                self._write(ws, r, 2, v["title"], fill=shade, wrap=True)
                self._write(ws, r, 3, v["views"], fill=shade, halign="center")
                self._write(ws, r, 4, v["likes"], fill=shade, halign="center")
                url = v["url"]
                if is_url(url):
                    self._write(ws, r, 5, "Watch ↗", fill=shade, link=url, halign="center")
                else:
                    self._write(ws, r, 5, "—", fill=shade, halign="center")

    def _sheet_transcripts(self, wb, d):
        sigs = d.transcript.get("signals", [])
        if not sigs:
            return
        ws = wb.create_sheet("Transcript Signals")
        ws.column_dimensions["A"].width = 10
        ws.column_dimensions["B"].width = 14
        ws.column_dimensions["C"].width = 10
        ws.column_dimensions["D"].width = 62
        ws.freeze_panes = "A3"

        self._section_banner(ws, 1,
            f"Internal Transcript · {d.transcript.get('source_file')} · "
            f"{d.transcript.get('total_signals')} signals · "
            f"{d.transcript.get('meeting_type')}", 4)
        self._header_row(ws, 2, ["Signal ID", "Type", "Confidence", "Content"])

        type_colours = {
            "Trend":       C_ACCENT_LIGHT,
            "Risk":        C_RED_LIGHT,
            "Opportunity": C_GREEN_LIGHT,
            "Feature":     "E8F5E9",
            "Pain Point":  C_AMBER_LIGHT,
        }
        for i, s in enumerate(sigs, 1):
            r = i + 2
            shade = type_colours.get(s["type"], C_GREY_BG if i % 2 == 0 else C_WHITE)
            conf = s.get("confidence")
            conf_str = f"{float(conf):.0%}" if conf is not None else "—"
            self._write(ws, r, 1, s["id"],      fill=shade, halign="center")
            self._write(ws, r, 2, s["type"],    fill=shade, bold=True)
            self._write(ws, r, 3, conf_str,     fill=shade, halign="center")
            self._write(ws, r, 4, s["content"], fill=shade, wrap=True)
            ws.row_dimensions[r].height = max(15, min(90, len(s["content"]) // 5 + 15))

    def _sheet_problems(self, wb, d):
        if not d.problems:
            return
        ws = wb.create_sheet("Problems")
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 46
        ws.column_dimensions["C"].width = 11
        ws.column_dimensions["D"].width = 11
        ws.column_dimensions["E"].width = 14
        ws.column_dimensions["F"].width = 12
        ws.column_dimensions["G"].width = 38
        ws.freeze_panes = "A3"

        self._section_banner(ws, 1,
            f"Identified Problems  —  Total: {d.total_problems}  |  "
            f"Top Categories: {d.top_categories}  |  "
            f"High Severity: {d.high_severity_count}", 7)
        self._header_row(ws, 2, ["ID", "Problem Statement", "Severity",
                                  "Frequency", "Category", "User Type", "Evidence Quotes"])
        for i, p in enumerate(d.problems, 1):
            r = i + 2
            shade = self._sev_fill(p["severity"])
            ev_str = "\n".join(f"• {e}" for e in p["evidence"]) if p["evidence"] else "—"
            self._write(ws, r, 1, p["id"],       fill=shade, halign="center", bold=True)
            self._write(ws, r, 2, p["problem"],  fill=shade, wrap=True)
            sev_txt = p["severity"]
            sev_fg  = SEV_PALETTE.get(sev_txt.lower(), (C_TEXT_DARK, C_WHITE))[0]
            c = ws.cell(row=r, column=3, value=sev_txt)
            c.font      = Font(name=FONT_NAME, size=10, bold=True, color=sev_fg)
            c.fill      = self._P(shade)
            c.alignment = self._A("center", "top", False)
            c.border    = self.BORDER
            self._write(ws, r, 4, p["frequency"],     fill=shade, halign="center")
            self._write(ws, r, 5, p["category"],      fill=shade)
            self._write(ws, r, 6, p["user_type"],     fill=shade)
            self._write(ws, r, 7, ev_str,             fill=shade, wrap=True, italic=True)
            ws.row_dimensions[r].height = max(20, min(100, len(p["problem"]) // 3 + 20))

    def _sheet_insights(self, wb, d):
        if not d.insights:
            return
        ws = wb.create_sheet("Strategic Insights")
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 44
        ws.column_dimensions["C"].width = 10
        ws.column_dimensions["D"].width = 12
        ws.column_dimensions["E"].width = 36
        ws.column_dimensions["F"].width = 36
        ws.freeze_panes = "A3"

        self._section_banner(ws, 1,
            f"Strategic Insights  —  Total: {d.total_insights}  |  "
            f"Dominant Theme: {d.dominant_theme}  |  "
            f"Critical: {d.critical_count}", 6)
        self._header_row(ws, 2, ["ID", "Insight", "Priority", "Confidence",
                                  "Root Cause", "Implication"])
        for i, ins in enumerate(d.insights, 1):
            r = i + 2
            shade = self._sev_fill(ins["priority"])
            self._write(ws, r, 1, ins["id"],         fill=shade, halign="center", bold=True)
            self._write(ws, r, 2, ins["insight"],    fill=shade, wrap=True)
            pri_fg = SEV_PALETTE.get(ins["priority"].lower(), (C_TEXT_DARK, C_WHITE))[0]
            c = ws.cell(row=r, column=3, value=ins["priority"])
            c.font      = Font(name=FONT_NAME, size=10, bold=True, color=pri_fg)
            c.fill      = self._P(shade)
            c.alignment = self._A("center", "top", False)
            c.border    = self.BORDER
            self._write(ws, r, 4, ins["confidence"], fill=shade, halign="center")
            self._write(ws, r, 5, ins["root_cause"], fill=shade, wrap=True)
            self._write(ws, r, 6, ins["implication"],fill=shade, wrap=True)
            ws.row_dimensions[r].height = max(20, min(100, len(ins["insight"]) // 3 + 20))

        # Strategic summary block
        r = len(d.insights) + 4
        self._section_banner(ws, r, "Strategic Summary", 6); r += 1
        for lbl, val in [("Key Strategic Risk",   d.strategic_risk),
                         ("Biggest Opportunity",  d.biggest_opp)]:
            if not is_empty(val):
                shade = C_GREY_BG if r % 2 == 0 else C_WHITE
                self._write(ws, r, 1, lbl, bold=True, fill=shade)
                ws.merge_cells(start_row=r, start_column=2,
                               end_row=r, end_column=6)
                self._write(ws, r, 2, val, fill=shade, wrap=True); r += 1

    def _sheet_briefs(self, wb, d):
        if not d.briefs:
            return
        ws = wb.create_sheet("Product Briefs")
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 9
        ws.column_dimensions["D"].width = 9
        ws.column_dimensions["E"].width = 36
        ws.column_dimensions["F"].width = 36
        ws.column_dimensions["G"].width = 30
        ws.freeze_panes = "A3"

        self._section_banner(ws, 1,
            f"Product Briefs  —  Total: {d.total_briefs}  |  "
            f"Sprint Focus: {d.sprint_focus[:80]+'…' if len(d.sprint_focus) > 80 else d.sprint_focus}", 7)
        self._header_row(ws, 2, ["ID", "Feature Name", "Priority", "Effort",
                                  "Problem → Solution", "Expected Impact", "User Flow"])
        for i, b in enumerate(d.briefs, 1):
            r = i + 2
            shade = self._sev_fill(b["priority"])
            prob_sol = f"PROBLEM:\n{b['problem']}\n\nSOLUTION:\n{b['solution']}"
            flow = "\n".join(f"{j+1}. {s}" for j, s in enumerate(b["user_flow"])) if b["user_flow"] else "—"
            self._write(ws, r, 1, b["id"],      fill=shade, halign="center", bold=True)
            self._write(ws, r, 2, b["feature"], fill=shade, wrap=True, bold=True)
            pri_fg = SEV_PALETTE.get(b["priority"].lower(), (C_TEXT_DARK, C_WHITE))[0]
            c = ws.cell(row=r, column=3, value=b["priority"])
            c.font      = Font(name=FONT_NAME, size=10, bold=True, color=pri_fg)
            c.fill      = self._P(shade)
            c.alignment = self._A("center", "top", False)
            c.border    = self.BORDER
            self._write(ws, r, 4, b["effort"],  fill=shade, halign="center")
            self._write(ws, r, 5, prob_sol,     fill=shade, wrap=True)
            self._write(ws, r, 6, b["impact"],  fill=shade, wrap=True)
            self._write(ws, r, 7, flow,         fill=shade, wrap=True)
            ws.row_dimensions[r].height = max(30, min(120, len(prob_sol) // 4 + 20))

    def _sheet_company_detail(self, wb, d):
        ws = wb.create_sheet("Company Detail")
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 60
        ws.freeze_panes = "A2"

        self._header_row(ws, 1, ["Attribute", "Detail"])

        def _block(title, rows):
            nonlocal r
            r += 1
            self._section_banner(ws, r, title, 2); r += 1
            for key, val in rows:
                if not is_empty(val):
                    self._kv_row(ws, r, key, val); r += 1

        r = 1
        _block("Milestones", [(f"Milestone {i+1}", m) for i, m in enumerate(d.milestones)])
        _block("New Features Launched", [(f"Feature {i+1}", f) for i, f in enumerate(d.new_features)])
        _block("Differentiators",
               [(item.get("feature", f"Item {i+1}"),
                 item.get("feature") or item.get("value", ""))
                for i, item in enumerate(d.differentiators)])
        _block("Competitors",
               [(item.get("name", f"#{i+1}"),
                 f"Domain: {item.get('domain','—')}")
                for i, item in enumerate(d.competitors)])
        _block("Strategic Moves",
               [(f"Move {i+1}", item.get("move") or item.get("value", ""))
                for i, item in enumerate(d.strategic_moves)])
        _block("Partnerships",
               [(item.get("partner", f"#{i+1}"),
                 f"{item.get('type','—')} · {item.get('description','—')} · {item.get('date','—')}")
                for i, item in enumerate(d.partnerships)])
        _block("User Complaints",
               [(f"Complaint {i+1}", item.get("issue") or item.get("value", ""))
                for i, item in enumerate(d.user_complaints)])
        _block("Regulatory & Legal",
               [(item.get("issue", f"Issue {i+1}"),
                 f"{item.get('status','—')} · {item.get('jurisdiction','—')} · {item.get('date','—')}")
                for i, item in enumerate(d.regulatory)])
        _block("Current Problems / Struggles",
               [(f"Problem {i+1}", item.get("description") or item.get("value",""))
                for i, item in enumerate(d.current_problems)])
        _block("Other Crucial Details", [(f"Detail {i+1}", t) for i, t in enumerate(d.other_details)])

    def build(self, d, logo_bytes, out_path):
        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        # Collect any temp logo file paths so we can clean up AFTER wb.save()
        logo_tmp_holder = []

        self._sheet_overview(wb, d, logo_bytes, _logo_tmp_holder=logo_tmp_holder)
        self._sheet_app_store(wb, d)
        self._sheet_reviews(wb, d)
        self._sheet_social(wb, d)
        self._sheet_transcripts(wb, d)
        self._sheet_problems(wb, d)
        self._sheet_insights(wb, d)
        self._sheet_briefs(wb, d)
        self._sheet_company_detail(wb, d)

        try:
            wb.save(out_path)
        finally:
            # Safe to delete temp files only after the workbook has been written
            for tmp_path in logo_tmp_holder:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
        print(f"  ✓ Excel  →  {out_path}")



# ══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER  (pure python-docx, no Node/npm required)
# ══════════════════════════════════════════════════════════════════════════════

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── palette (same as Excel) ───────────────────────────────────────────────────
def _rgb(hex6):
    h = hex6.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

DC   = _rgb("0D1F3C")   # dark navy
ACC  = _rgb("1A6BCC")   # corporate blue
ACCL = _rgb("D6E8FF")   # pale blue
GRN  = _rgb("1A7A4A")
GRNL = _rgb("D4EDDA")
AMB  = _rgb("B76E00")
AMBL = _rgb("FFF3CD")
RED  = _rgb("B71C1C")
REDL = _rgb("FDECEA")
GRY  = _rgb("F5F7FA")   # alt row
LIN  = _rgb("DEE2E8")   # border
WHT  = _rgb("FFFFFF")
MUT  = _rgb("6B7280")

SEV_COL = {
    "critical": (RED,  REDL),
    "high":     (AMB,  AMBL),
    "medium":   (_rgb("7B5800"), _rgb("FFFBEA")),
    "low":      (GRN,  GRNL),
    "p0":       (RED,  REDL),
    "p1":       (AMB,  AMBL),
    "p2":       (_rgb("7B5800"), _rgb("FFFBEA")),
    "p3":       (GRN,  GRNL),
}
def sev_col(s):
    return SEV_COL.get(str(s).lower(), (DC, WHT))

TYPE_COL = {
    "Trend":       ACC,
    "Risk":        RED,
    "Opportunity": GRN,
    "Feature":     _rgb("2C7A4B"),
    "Pain Point":  AMB,
}

FONT = "Calibri"

# ── low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex6):
    hex6 = hex6.lstrip("#")
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6.upper())
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def _set_cell_borders(cell, color_hex="DEE2E8", size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","bottom","left","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex.upper())
        borders.append(el)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)

def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(mar)

def _set_col_width(cell, width_dxa):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.insert(0, tcW)

def _set_table_width(table, width_dxa=9360):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)

def _set_para_border_left(para, color_hex, size=20):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex.lstrip("#").upper())
    pBdr.append(left)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)

def _set_para_bottom_border(para, color_hex, size=8):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex.lstrip("#").upper())
    pBdr.append(bot)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)

def _set_para_bg(para, hex6):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6.lstrip("#").upper())
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)

def _set_para_indent(para, left=0, right=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    if left:  ind.set(qn("w:left"),  str(left))
    if right: ind.set(qn("w:right"), str(right))
    existing = pPr.find(qn("w:ind"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(ind)

def _set_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spc)

def _add_hyperlink(para, text, url):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl   = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r    = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    style= OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rPr.append(style)
    # colour + underline
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1A6BCC")
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(color_el); rPr.append(u_el)
    rPr.append(OxmlElement("w:rFonts"))
    rPr[-1].set(qn("w:ascii"), FONT)
    rPr[-1].set(qn("w:hAnsi"), FONT)
    t    = OxmlElement("w:t")
    t.text = text
    r.append(rPr); r.append(t)
    hl.append(r)
    para._p.append(hl)

# ── high-level helpers ────────────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size=10, color=None, font=FONT):
    run = para.add_run(str(text) if text is not None else "")
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def _heading(doc, text, level=1, color=None, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    colors= {1: DC,  2: ACC, 3: _rgb("2C3E50")}
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    _set_spacing(p, before=240 if level==1 else 180, after=120 if level==1 else 80)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(size or sizes[level])
    r.font.color.rgb = color or colors[level]
    return p

def _blank(doc):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=60)
    return p

def _bullet(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
    _run(p, text, size=10)
    _set_spacing(p, before=40, after=40)
    return p

def _is_url(v):
    return isinstance(v, str) and re.match(r"https?://", v.strip())

def _nonempty(v):
    return v not in (None, "", "—", "—", "null", "none", "n/a", "na",
                     "unable to verify", "not available", "not found")

# ── table builders ────────────────────────────────────────────────────────────

def _prep_cell(cell, bg_hex, width_dxa=None):
    _set_cell_bg(cell, bg_hex.lstrip("#"))
    _set_cell_borders(cell)
    _cell_margins(cell)
    if width_dxa:
        _set_col_width(cell, width_dxa)
    # clear default paragraph
    for p in cell.paragraphs:
        _set_spacing(p, 0, 0)

def _hdr_cell(cell, text, width_dxa=None):
    _prep_cell(cell, "0D1F3C", width_dxa)
    p = cell.paragraphs[0]
    _run(p, text, bold=True, color=WHT, size=10)

def _str_cell(cell, text, bg_hex="FFFFFF", width_dxa=None, link_url=None):
    _prep_cell(cell, bg_hex, width_dxa)
    p = cell.paragraphs[0]
    if link_url and _is_url(link_url):
        _add_hyperlink(p, str(text), link_url)
    else:
        _run(p, text, size=10)

def _bold_cell(cell, text, bg_hex="D6E8FF", width_dxa=None):
    _prep_cell(cell, bg_hex, width_dxa)
    p = cell.paragraphs[0]
    _run(p, text, bold=True, size=10)

def _kv_table(doc, pairs, col_widths=(2640, 6720)):
    """Two-column key/value table. pairs = [(label, value), ...] or [(label, value, is_link), ...]"""
    clean = []
    for p in pairs:
        k = p[0]; v = p[1]; as_link = p[2] if len(p) > 2 else False
        if _nonempty(v) and v != "—":
            clean.append((k, v, as_link))
    if not clean:
        return
    table = doc.add_table(rows=len(clean), cols=2)
    _set_table_width(table)
    for i, (k, v, as_link) in enumerate(clean):
        row  = table.rows[i]
        key_bg  = "E8F0FC" if i % 2 == 0 else "D6E8FF"
        val_bg  = "F5F7FA" if i % 2 == 0 else "FFFFFF"
        _bold_cell(row.cells[0], k, key_bg, col_widths[0])
        if as_link and _is_url(str(v)):
            _str_cell(row.cells[1], str(v), val_bg, col_widths[1], link_url=str(v))
        else:
            _str_cell(row.cells[1], str(v), val_bg, col_widths[1])
    return table

def _callout(doc, label, text, bg_hex, border_hex, label_color=None, text_color=None):
    p = doc.add_paragraph()
    _set_para_bg(p, bg_hex)
    _set_para_border_left(p, border_hex, size=24)
    _set_para_indent(p, left=240, right=240)
    _set_spacing(p, before=120, after=120)
    _run(p, label, bold=True, color=label_color or _rgb(border_hex.lstrip("#")), size=10)
    _run(p, text,  italic=True, color=text_color or DC, size=10)
    return p

def _section_divider(doc, label):
    p = doc.add_paragraph()
    _set_para_bottom_border(p, "1A6BCC", size=8)
    _set_spacing(p, before=200, after=80)
    _run(p, label, bold=True, color=ACC, size=12)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# MAIN DOCX BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_docx(d: "ReportData", logo_bytes: bytes, out_path: str):
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── COVER ────────────────────────────────────────────────────────────────
    if logo_bytes:
        try:
            buf = io.BytesIO(logo_bytes)
            p   = doc.add_paragraph()
            _set_spacing(p, before=0, after=120)
            run = p.add_run()
            run.add_picture(buf, width=Inches(0.75))
        except Exception:
            pass

    p = doc.add_paragraph()
    _set_spacing(p, before=60, after=60)
    _run(p, d.company_name, bold=True, size=26, color=DC)

    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=60)
    _run(p, clean_domain(d.company_domain), size=13, color=ACC)

    # Divider line
    div = doc.add_paragraph()
    _set_para_bottom_border(div, "1A6BCC", size=12)
    _set_spacing(div, before=0, after=240)

    # Cover summary 2-col table
    cover_kv = [
        ("Industry",      d.industry),
        ("Founded",       d.year_founded),
        ("Headquarters",  d.hq_location),
        ("Employees",     d.employee_count),
        ("Revenue",       d.annual_revenue),
        ("Funding",       f"{safe(d.funding_stage)} · {safe(d.funding_raised)}"),
        ("Platforms",     d.platforms),
        ("Revenue Model", d.revenue_model),
    ]
    cover_kv = [(k, v) for k, v in cover_kv if _nonempty(v) and "— · —" not in str(v)]
    if cover_kv:
        half   = (len(cover_kv) + 1) // 2
        left_  = cover_kv[:half]
        right_ = cover_kv[half:]
        maxr   = max(len(left_), len(right_))
        table  = doc.add_table(rows=maxr, cols=4)
        _set_table_width(table, 9360)
        for i in range(maxr):
            row = table.rows[i]
            lk, lv = left_[i]  if i < len(left_)  else ("", "")
            rk, rv = right_[i] if i < len(right_) else ("", "")
            _bold_cell(row.cells[0], lk, "E8F0FC" if i%2==0 else "D6E8FF", 2160)
            _str_cell (row.cells[1], lv, "F5F7FA" if i%2==0 else "FFFFFF",  2520)
            _bold_cell(row.cells[2], rk, "E8F0FC" if i%2==0 else "D6E8FF", 2160)
            _str_cell (row.cells[3], rv, "F5F7FA" if i%2==0 else "FFFFFF",  2520)
        _blank(doc)

    if _nonempty(d.key_positioning):
        _callout(doc, "Key Positioning:  ", d.key_positioning, "D6E8FF", "1A6BCC")
        _blank(doc)

    doc.add_page_break()

    # ── 1. COMPANY PROFILE ───────────────────────────────────────────────────
    _heading(doc, "1. Company Profile", level=1)
    _blank(doc)

    links = [(lbl, url) for lbl, url in [
        ("Play Store", d.playstore_link), ("App Store", d.appstore_link),
        ("YouTube",    d.youtube_channel), ("LinkedIn",  d.linkedin_page),
    ] if _is_url(url)]
    if links:
        _heading(doc, "Official Links", level=2)
        table = doc.add_table(rows=len(links), cols=2)
        _set_table_width(table)
        for i, (lbl, url) in enumerate(links):
            _bold_cell(table.rows[i].cells[0], lbl, "E8F0FC" if i%2==0 else "D6E8FF", 1800)
            _str_cell (table.rows[i].cells[1], url, "F5F7FA" if i%2==0 else "FFFFFF", 7560, link_url=url)
        _blank(doc)

    csuite  = [x for x in d.csuite   if _nonempty(x)]
    founders= [x for x in d.founders if _nonempty(x)]
    if csuite or founders:
        _heading(doc, "Leadership", level=2)
        rows_data = [("Executive", x) for x in csuite]
        if founders:
            rows_data.append(("Founders", "  •  ".join(founders)))
        table = doc.add_table(rows=len(rows_data), cols=2)
        _set_table_width(table)
        for i, (k, v) in enumerate(rows_data):
            _bold_cell(table.rows[i].cells[0], k, "E8F0FC" if i%2==0 else "D6E8FF", 1800)
            _str_cell (table.rows[i].cells[1], v, "F5F7FA" if i%2==0 else "FFFFFF",  7560)
        _blank(doc)

    if d.pricing_tiers:
        _heading(doc, "Pricing Tiers", level=2)
        for t in d.pricing_tiers:
            if _nonempty(t): _bullet(doc, t)
        _blank(doc)

    if d.target_segments:
        _heading(doc, "Target Customer Segments", level=2)
        for s in d.target_segments:
            if _nonempty(s): _bullet(doc, s)
        _blank(doc)

    if d.tech_stack:
        _heading(doc, "Technology Highlights", level=2)
        for t in d.tech_stack:
            if _nonempty(t): _bullet(doc, t)
        _blank(doc)

    ms = d.market_sentiment if isinstance(d.market_sentiment, dict) else {}
    if ms.get("overall") or ms.get("analyst_view"):
        _heading(doc, "Market Sentiment", level=2)
        _kv_table(doc, [
            ("Overall",        ms.get("overall")),
            ("Analyst View",   ms.get("analyst_view")),
            ("Community View", ms.get("user_community_view")),
            ("As of",          ms.get("date")),
        ])
        _blank(doc)

    doc.add_page_break()

    # ── 2. APP STORE ANALYSIS ────────────────────────────────────────────────
    ps  = d.ps  or {}
    app = d.app or {}
    if ps or app:
        _heading(doc, "2. App Store Analysis", level=1)
        _blank(doc)

        comparisons = [
            ("App Title",      ps.get("title"),           app.get("title")),
            ("Rating",         f"{float(ps['score']):.1f} ★" if ps.get("score") else "—",
                               f"{float(app['score']):.1f} ★" if app.get("score") else "—"),
            ("Total Ratings",  ps.get("ratings"),         app.get("ratings")),
            ("Installs",       ps.get("installs"),        "—"),
            ("Genre",          ps.get("genre"),           app.get("genre")),
            ("Price",          ps.get("free"),            app.get("price")),
            ("Released",       ps.get("released"),        app.get("released")),
            ("Version",        ps.get("version"),         app.get("version")),
            ("Content Rating", ps.get("content_rating"),  app.get("content_rating")),
        ]
        comparisons = [(l, a, b) for l, a, b in comparisons if _nonempty(a) or _nonempty(b)]
        if comparisons:
            _heading(doc, "Play Store vs App Store", level=2)
            table = doc.add_table(rows=len(comparisons)+1, cols=3)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Metric",            2400)
            _hdr_cell(table.rows[0].cells[1], "Google Play",       3480)
            _hdr_cell(table.rows[0].cells[2], "Apple App Store",   3480)
            for i, (lbl, pv, av) in enumerate(comparisons, 1):
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                _bold_cell(table.rows[i].cells[0], lbl, "E8F0FC" if i%2==0 else "D6E8FF", 2400)
                _str_cell (table.rows[i].cells[1], safe(pv), shade, 3480)
                _str_cell (table.rows[i].cells[2], safe(av), shade, 3480)
            _blank(doc)

        rd = ps.get("rating_dist", {})
        rd_entries = [(k, v) for k, v in rd.items() if _nonempty(v)]
        if rd_entries:
            _heading(doc, "Play Store Rating Distribution", level=2)
            total = sum(int(v or 0) for _, v in rd_entries) or 1
            table = doc.add_table(rows=6, cols=3)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Stars",        1440)
            _hdr_cell(table.rows[0].cells[1], "Count",        1440)
            _hdr_cell(table.rows[0].cells[2], "Distribution", 6480)
            for i, star in enumerate(["5","4","3","2","1"], 1):
                cnt  = int(rd.get(star, rd.get(int(star), 0)) or 0)
                bar  = "█" * int((cnt/total)*30)
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                bar_col = RED if star in ("1","2") else GRN
                _str_cell(table.rows[i].cells[0], f"{'⭐'*int(star)} {star}★", shade, 1440)
                _str_cell(table.rows[i].cells[1], str(cnt), shade, 1440)
                p = table.rows[i].cells[2].paragraphs[0]
                _prep_cell(table.rows[i].cells[2], shade.lstrip("#"), 6480)
                _run(p, bar or "—", color=bar_col, size=10)
            _blank(doc)

        reviews = ps.get("reviews_list", [])
        if reviews:
            _heading(doc, f"Top Reviews ({len(reviews)} analysed)", level=2)
            sorted_rev = sorted(reviews, key=lambda r: int(r.get("rating", 5) or 5))
            for rev in sorted_rev[:8]:
                sc  = int(rev.get("rating", 3) or 3)
                sev = "critical" if sc <= 2 else ("medium" if sc <= 3 else "low")
                tc_, bgc = sev_col(sev)
                bg_hex = ("FDECEA" if sev=="critical" else ("FFFBEA" if sev=="medium" else "D4EDDA"))
                tc_hex_map = {"critical": "B71C1C", "medium": "7B5800", "low": "1A7A4A"}
                tc_hex = tc_hex_map.get(sev, "B71C1C")
                p = doc.add_paragraph()
                _set_para_bg(p, bg_hex)
                _set_para_border_left(p, tc_hex, size=20)
                _set_para_indent(p, left=240)
                _set_spacing(p, before=100, after=0)
                _run(p, f"{'⭐'*sc}  {safe(rev.get('author'))}",
                     bold=True, color=tc_, size=10)
                _run(p, f"  ·  {safe(rev.get('date'))}", color=MUT, size=9)

                p2 = doc.add_paragraph()
                _set_para_bg(p2, bg_hex)
                _set_para_indent(p2, left=240, right=240)
                _set_spacing(p2, before=0, after=0 if rev.get("reply") else 120)
                _run(p2, safe(rev.get("content")), size=10)

                if rev.get("reply") and _nonempty(rev["reply"]):
                    p3 = doc.add_paragraph()
                    _set_para_bg(p3, "F0F4FF")
                    _set_para_indent(p3, left=360, right=240)
                    _set_spacing(p3, before=0, after=120)
                    _run(p3, "Developer: ", bold=True, color=ACC, size=9)
                    _run(p3, safe(rev["reply"]), italic=True, color=MUT, size=9)
            _blank(doc)
        doc.add_page_break()

    # ── 3. SOCIAL & MEDIA ────────────────────────────────────────────────────
    reddit  = d.reddit_posts   or []
    youtube = d.youtube_videos or []
    if reddit or youtube:
        _heading(doc, "3. Social & Media Intelligence", level=1)
        _blank(doc)

        if reddit:
            _heading(doc, f"Reddit Posts  ({len(reddit)})", level=2)
            table = doc.add_table(rows=len(reddit)+1, cols=4)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Post Title",  5640)
            _hdr_cell(table.rows[0].cells[1], "Subreddit",   1080)
            _hdr_cell(table.rows[0].cells[2], "Score",       1080)
            _hdr_cell(table.rows[0].cells[3], "Link",        1560)
            for i, p in enumerate(reddit, 1):
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                _str_cell(table.rows[i].cells[0], safe(p.get("title")),         shade, 5640)
                _str_cell(table.rows[i].cells[1], f"r/{safe(p.get('subreddit'))}", shade, 1080)
                _str_cell(table.rows[i].cells[2], safe(p.get("score")),         shade, 1080)
                url = safe(p.get("url"))
                if _is_url(url):
                    _str_cell(table.rows[i].cells[3], "Open ↗", shade, 1560, link_url=url)
                else:
                    _str_cell(table.rows[i].cells[3], "—", shade, 1560)
            _blank(doc)

        if youtube:
            _heading(doc, f"YouTube Coverage  ({len(youtube)} videos)", level=2)
            table = doc.add_table(rows=len(youtube)+1, cols=4)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Video Title", 6120)
            _hdr_cell(table.rows[0].cells[1], "Views",       1080)
            _hdr_cell(table.rows[0].cells[2], "Likes",       1080)
            _hdr_cell(table.rows[0].cells[3], "Link",        1080)
            for i, v in enumerate(youtube, 1):
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                _str_cell(table.rows[i].cells[0], safe(v.get("title")), shade, 6120)
                _str_cell(table.rows[i].cells[1], safe(v.get("views")), shade, 1080)
                _str_cell(table.rows[i].cells[2], safe(v.get("likes")), shade, 1080)
                url = safe(v.get("url"))
                if _is_url(url):
                    _str_cell(table.rows[i].cells[3], "Watch ↗", shade, 1080, link_url=url)
                else:
                    _str_cell(table.rows[i].cells[3], "—", shade, 1080)
            _blank(doc)
        doc.add_page_break()

    # ── 4. TRANSCRIPT SIGNALS ────────────────────────────────────────────────
    signals = d.transcript.get("signals", []) if d.transcript else []
    if signals:
        _heading(doc, "4. Internal Transcript Signals", level=1)
        _blank(doc)
        tr = d.transcript
        _kv_table(doc, [
            ("Source File",   tr.get("source_file")),
            ("Total Signals", tr.get("total_signals")),
            ("Classifier",    tr.get("classifier")),
            ("Meeting Type",  tr.get("meeting_type")),
        ])
        _blank(doc)
        TYPE_HEX = {
            "Trend": "1A6BCC", "Risk": "B71C1C", "Opportunity": "1A7A4A",
            "Feature": "2C7A4B", "Pain Point": "B76E00",
        }
        for sig in signals:
            col_hex = TYPE_HEX.get(sig.get("type"), "0D1F3C")
            col     = _rgb(col_hex)
            conf = sig.get("confidence")
            conf_str = f"{float(conf)*100:.0f}%" if conf is not None else "—"
            p = doc.add_paragraph()
            _set_para_bg(p, "F4F6FA")
            _set_para_border_left(p, col_hex, size=16)
            _set_para_indent(p, left=240)
            _set_spacing(p, before=100, after=0)
            _run(p, f"[{safe(sig.get('id'))}]  {safe(sig.get('type'))}", bold=True, color=col, size=10)
            _run(p, f"  ·  Confidence: {conf_str}", color=MUT, size=9)
            p2 = doc.add_paragraph()
            _set_para_bg(p2, "F4F6FA")
            _set_para_indent(p2, left=240, right=240)
            _set_spacing(p2, before=0, after=120)
            _run(p2, safe(sig.get("content")), size=9)
        doc.add_page_break()

    # ── 5. IDENTIFIED PROBLEMS ───────────────────────────────────────────────
    problems = d.problems or []
    if problems:
        _heading(doc, "5. Identified Problems", level=1)
        _blank(doc)

        # Summary stats row
        table = doc.add_table(rows=1, cols=4)
        _set_table_width(table)
        stats = [
            ("Total Problems",  safe(d.total_problems or len(problems)), "0D1F3C", "FFFFFF"),
            ("Critical / High", safe(d.high_severity_count or "—"),      "B71C1C", "FDECEA"),
            ("Top Categories",  safe(d.top_categories),                   "1A6BCC", "D6E8FF"),
            ("Sources Used",    safe(getattr(d, "sources_used", "—")),   "1A7A4A", "D4EDDA"),
        ]
        for ci, (label, val, fg, bg) in enumerate(stats):
            cell = table.rows[0].cells[ci]
            _prep_cell(cell, bg, 2340)
            p1 = cell.paragraphs[0]
            _run(p1, label, bold=True, color=_rgb(fg), size=9)
            p2 = cell.add_paragraph()
            _set_spacing(p2, before=40, after=0)
            _run(p2, val,   bold=True, color=_rgb(fg), size=16)
        _blank(doc)

        SEV_HEX = {"critical":"B71C1C","p0":"B71C1C","high":"B76E00","p1":"B76E00",
                   "medium":"7B5800","p2":"7B5800","low":"1A7A4A","p3":"1A7A4A"}
        SEV_BG  = {"critical":"FDECEA","p0":"FDECEA","high":"FFF3CD","p1":"FFF3CD",
                   "medium":"FFFBEA","p2":"FFFBEA","low":"D4EDDA","p3":"D4EDDA"}
        for prob in problems:
            sev_key = str(prob.get("severity","")).lower()
            tc_hex  = SEV_HEX.get(sev_key, "0D1F3C")
            bg_hex  = SEV_BG.get(sev_key, "F5F7FA")
            tc_     = _rgb(tc_hex)

            p = doc.add_paragraph()
            _set_spacing(p, before=200, after=0)
            _run(p, f"{safe(prob.get('id'))} — ", bold=True, color=DC, size=11)
            _run(p, safe(prob.get("severity","")).upper(), bold=True, color=tc_, size=10)
            _run(p, f"  ·  {safe(prob.get('category'))}  ·  {safe(prob.get('frequency'))}", color=MUT, size=9)

            p2 = doc.add_paragraph()
            _set_para_bg(p2, bg_hex)
            _set_para_border_left(p2, tc_hex, size=20)
            _set_para_indent(p2, left=240, right=160)
            _set_spacing(p2, before=60, after=80)
            _run(p2, safe(prob.get("problem")), size=10)

            evidence = [e for e in (prob.get("evidence") or []) if _nonempty(e)]
            if evidence:
                ep = doc.add_paragraph()
                _set_spacing(ep, before=40, after=20)
                _run(ep, "Evidence:", bold=True, color=DC, size=9)
                for ev in evidence:
                    _bullet(doc, f'"{ev}"')
            _blank(doc)
        doc.add_page_break()

    # ── 6. STRATEGIC INSIGHTS ────────────────────────────────────────────────
    insights = d.insights or []
    if insights:
        _heading(doc, "6. Strategic Insights", level=1)
        _blank(doc)

        if _nonempty(d.strategic_risk):
            _callout(doc, "⚠  Strategic Risk: ", d.strategic_risk, "FDECEA", "B71C1C")
        if _nonempty(d.biggest_opp):
            _callout(doc, "✓  Opportunity: ", d.biggest_opp, "D4EDDA", "1A7A4A")
        _blank(doc)

        SEV_HEX = {"critical":"B71C1C","p0":"B71C1C","high":"B76E00","p1":"B76E00",
                   "medium":"7B5800","p2":"7B5800","low":"1A7A4A","p3":"1A7A4A"}
        SEV_BG  = {"critical":"FDECEA","p0":"FDECEA","high":"FFF3CD","p1":"FFF3CD",
                   "medium":"FFFBEA","p2":"FFFBEA","low":"D4EDDA","p3":"D4EDDA"}
        for ins in insights:
            sev_key = str(ins.get("priority","")).lower()
            tc_hex  = SEV_HEX.get(sev_key, "0D1F3C")
            bg_hex  = SEV_BG.get(sev_key, "F5F7FA")
            tc_     = _rgb(tc_hex)

            p = doc.add_paragraph()
            _set_spacing(p, before=200, after=0)
            _run(p, f"{safe(ins.get('id'))} — ", bold=True, color=DC, size=11)
            _run(p, safe(ins.get("theme","")).upper(), bold=True, color=tc_, size=10)
            _run(p, f"  ·  {safe(ins.get('priority'))}  ·  Confidence: {safe(ins.get('confidence'))}", color=MUT, size=9)

            p2 = doc.add_paragraph()
            _set_para_bg(p2, bg_hex)
            _set_para_border_left(p2, tc_hex, size=20)
            _set_para_indent(p2, left=240, right=160)
            _set_spacing(p2, before=60, after=80)
            _run(p2, safe(ins.get("insight")), size=10)

            _kv_table(doc, [
                ("Root Cause",     ins.get("root_cause")),
                ("Evidence",       ins.get("evidence")),
                ("Competitor Gap", ins.get("competitor_gap")),
                ("Opportunity",    ins.get("opportunity")),
                ("Implication",    ins.get("implication")),
            ])
            _blank(doc)
        doc.add_page_break()

    # ── 7. PRODUCT BRIEFS ────────────────────────────────────────────────────
    briefs = d.briefs or []
    if briefs:
        _heading(doc, "7. Product Briefs", level=1)
        _blank(doc)

        if _nonempty(d.sprint_focus):
            _callout(doc, "Recommended Sprint Focus: ", d.sprint_focus, "D6E8FF", "1A6BCC")
            _blank(doc)

        for b in briefs:
            tc_, bgc = sev_col(b.get("priority",""))
            p = doc.add_paragraph()
            _set_spacing(p, before=200, after=0)
            _run(p, f"{safe(b.get('brief_id', b.get('id','')))}  ", bold=True, color=DC, size=11)
            _run(p, safe(b.get("feature_name", b.get("feature",""))), bold=True, color=ACC, size=11)
            _run(p, f"  [{safe(b.get('priority'))}]  ·  Effort: {safe(b.get('effort'))}", color=MUT, size=9)

            _kv_table(doc, [
                ("Problem",   b.get("problem")),
                ("Why Now",   b.get("why_now")),
                ("Solution",  b.get("solution")),
                ("Impact",    b.get("expected_impact", b.get("impact"))),
                ("Metric",    b.get("success_metric",  b.get("metric"))),
            ])

            uf = [s for s in (b.get("user_flow") or []) if _nonempty(s)]
            if uf:
                _blank(doc)
                ep = doc.add_paragraph()
                _set_spacing(ep, before=40, after=20)
                _run(ep, "User Flow:", bold=True, color=DC, size=9)
                for step in uf:
                    _bullet(doc, step, numbered=True)
            _blank(doc)
        doc.add_page_break()

    # ── 8. TIMELINE & STRATEGY ───────────────────────────────────────────────
    milestones = d.milestones or []
    moves      = d.strategic_moves or []
    partners   = d.partnerships or []
    if milestones or moves or partners:
        _heading(doc, "8. Timeline & Strategy", level=1)
        _blank(doc)

        if milestones:
            _heading(doc, "Key Milestones", level=2)
            for m in milestones:
                val = m.get("value") or m.get("milestone") or str(m) if isinstance(m, dict) else str(m)
                if _nonempty(val): _bullet(doc, val)
            _blank(doc)

        if moves:
            _heading(doc, "Strategic Moves", level=2)
            for m in moves:
                val = m.get("move") or m.get("value") or str(m) if isinstance(m, dict) else str(m)
                if _nonempty(val): _bullet(doc, val)
            _blank(doc)

        if partners:
            _heading(doc, "Partnerships & Integrations", level=2)
            table = doc.add_table(rows=len(partners)+1, cols=4)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Partner",     2160)
            _hdr_cell(table.rows[0].cells[1], "Type",        1440)
            _hdr_cell(table.rows[0].cells[2], "Description", 4320)
            _hdr_cell(table.rows[0].cells[3], "Date",        1440)
            for i, pt in enumerate(partners, 1):
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                _bold_cell(table.rows[i].cells[0], safe(pt.get("partner", pt.get("value",""))), "E8F0FC" if i%2==0 else "D6E8FF", 2160)
                _str_cell (table.rows[i].cells[1], safe(pt.get("type")),        shade, 1440)
                _str_cell (table.rows[i].cells[2], safe(pt.get("description")), shade, 4320)
                _str_cell (table.rows[i].cells[3], safe(pt.get("date")),        shade, 1440)
            _blank(doc)
        doc.add_page_break()

    # ── 9. COMPETITIVE LANDSCAPE ─────────────────────────────────────────────
    competitors   = d.competitors or []
    differentiators = d.differentiators or []
    if competitors or differentiators:
        _heading(doc, "9. Competitive Landscape", level=1)
        _blank(doc)

        if competitors:
            _heading(doc, "Competitors", level=2)
            table = doc.add_table(rows=len(competitors)+1, cols=2)
            _set_table_width(table)
            _hdr_cell(table.rows[0].cells[0], "Competitor",       3240)
            _hdr_cell(table.rows[0].cells[1], "Domain / Notes",   6120)
            for i, c in enumerate(competitors, 1):
                shade = "F5F7FA" if i%2==0 else "FFFFFF"
                name  = safe(c.get("name", c.get("value","")))
                dom   = safe(c.get("domain",""))
                _bold_cell(table.rows[i].cells[0], name, "E8F0FC" if i%2==0 else "D6E8FF", 3240)
                if _is_url(dom):
                    _str_cell(table.rows[i].cells[1], dom, shade, 6120, link_url=dom)
                else:
                    _str_cell(table.rows[i].cells[1], dom, shade, 6120)
            _blank(doc)

        if differentiators:
            _heading(doc, "Differentiators", level=2)
            for df in differentiators:
                val = df.get("feature") or df.get("value") or str(df) if isinstance(df, dict) else str(df)
                if _nonempty(val): _bullet(doc, val)
            _blank(doc)

    # Save
    doc.save(out_path)
    print(f"  ✓ DOCX   →  {out_path}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def main(json_path: str):
    print(f"\n{'='*60}")
    print(f"  Report Generator")
    print(f"  Input : {json_path}")
    print(f"{'='*60}\n")

    print("  Parsing JSON …")
    d = ReportData(json_path)
    print(f"  Company  : {d.company_name}")
    print(f"  Domain   : {clean_domain(d.company_domain)}")

    print("  Fetching logo …")
    logo_bytes, _ = fetch_logo_bytes(d.company_domain, d.company_name)
    print(f"  Logo     : {len(logo_bytes)} bytes")

    out_dir = Path("output") / slug(d.company_name)
    out_dir.mkdir(parents=True, exist_ok=True)
    base = slug(d.company_name)

    print("  Building Excel …")
    ExcelBuilder().build(d, logo_bytes, str(out_dir / f"{base}_report.xlsx"))

    print("  Building DOCX …")
    build_docx(d, logo_bytes, str(out_dir / f"{base}_report.docx"))

    print(f"\n  Done!  Output folder: {out_dir.resolve()}\n")


if __name__ == "__main__":
    json_file = sys.argv[1] if len(sys.argv) > 1 else "db_document.json"
    main(json_file)